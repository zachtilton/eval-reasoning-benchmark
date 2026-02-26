"""
Build Appendices G and H as a Word document.
Run: python3 appendices/build_appendices_gh.py
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

GITHUB_BASE = "https://github.com/zachtilton/eval-reasoning-benchmark/blob/main"

doc = Document()

# ── Global style tweaks ───────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

def set_margins(document, top=1, bottom=1, left=1.25, right=1.25):
    for section in document.sections:
        section.top_margin    = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin   = Inches(left)
        section.right_margin  = Inches(right)

set_margins(doc)

# ── Helper: add a hyperlink inline ───────────────────────────────────────────
def add_hyperlink(paragraph, text, url):
    """Insert a clickable hyperlink into an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Underline + blue colour (classic hyperlink style)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_hyperlink_run(paragraph, text, url):
    """Convenience wrapper that adds the hyperlink and returns the paragraph."""
    add_hyperlink(paragraph, text, url)
    return paragraph


# ── Helper: heading styles ────────────────────────────────────────────────────
def h1(text):
    """Appendix-level heading (e.g. 'Appendix G …')."""
    p = doc.add_heading(text, level=1)
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(14)
    p.runs[0].font.bold = True
    return p

def h2(text):
    """Sub-section heading (e.g. 'G.1 …')."""
    p = doc.add_heading(text, level=2)
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.bold = True
    return p

def h3(text):
    """Sub-sub-section heading."""
    p = doc.add_heading(text, level=3)
    p.runs[0].font.name = "Times New Roman"
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.bold = True
    return p


def body(text):
    """Normal body paragraph."""
    p = doc.add_paragraph(text)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(6)
    return p


def body_with_link(prefix, link_text, url, suffix=""):
    """Body paragraph containing an inline hyperlink."""
    p = doc.add_paragraph(prefix)
    p.style = doc.styles["Normal"]
    p.paragraph_format.space_after = Pt(6)
    add_hyperlink(p, link_text, url)
    if suffix:
        p.add_run(suffix)
    return p


# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX G
# ─────────────────────────────────────────────────────────────────────────────

h1("Appendix G: API Execution Infrastructure")

body(
    "This appendix documents the software infrastructure used to collect the 5,400 model "
    "responses that constitute the primary dataset of this study. The infrastructure is "
    "organized into five modules, each responsible for a distinct stage of the data "
    "collection or scoring pipeline. Subsections G.1 through G.3 describe the execution "
    "layer that configures, dispatches, and captures API calls; subsections G.4 and G.5 "
    "describe the automated scoring layer that validates, compares, and adjudicates those "
    "responses against the gold standard."
)

# ── G.1 ──────────────────────────────────────────────────────────────────────
h2("G.1  API Configuration and Batch Orchestration")

body(
    "This module establishes the technical foundation for the entire data collection phase. "
    "It defines how the benchmark is configured to interact with six separate artificial "
    "intelligence providers, how the 5,400 required API calls are planned and scheduled "
    "across multiple sessions, and how progress is monitored throughout execution. The "
    "relevant source files are "
)

p = doc.add_paragraph()
p.style = doc.styles["Normal"]
p.paragraph_format.space_after = Pt(6)
add_hyperlink(p, "config/api_config.py", f"{GITHUB_BASE}/config/api_config.py")
p.add_run(", ")
add_hyperlink(p, "config/model_params.py", f"{GITHUB_BASE}/config/model_params.py")
p.add_run(", and ")
add_hyperlink(p, "src/api_client/batch.py", f"{GITHUB_BASE}/src/api_client/batch.py")
p.add_run(".")

h3("Configuration Architecture")

body(
    "The benchmark interacts with six distinct language models maintained by different "
    "commercial providers: GPT 5.2 (OpenAI), Claude Opus 4.6 (Anthropic), Gemini 3 Pro "
    "(Google), DeepSeek V3 (DeepSeek), Kimi K2 (Moonshot AI), and GLM 4.7 (Zhipu AI). "
    "Each provider maintains its own web-based interface through which external software "
    "can request model responses. The configuration file serves as the authoritative "
    "registry for how to reach each of these interfaces, specifying the web address, the "
    "model identifier string that each provider uses internally, and the authentication "
    "protocol required to prove that requests originate from an authorized account. Three "
    "authentication mechanisms are employed across the six providers: a bearer token "
    "passed as a request header (OpenAI, DeepSeek, Kimi, and Zhipu), a specialized header "
    "format used by Anthropic, and a URL-embedded query parameter used by Google's Gemini. "
    "All API authentication credentials are stored as system environment variables rather "
    "than in the code itself, protecting sensitive keys from accidental exposure."
)

body(
    "A complementary configuration file defines the standardized generation parameters "
    "applied uniformly across all six models. Temperature is set to zero, which minimizes "
    "randomness in model outputs and maximizes determinism—the same prompt submitted twice "
    "to a temperature-zero model will reliably produce the same response. Maximum output "
    "length is capped at 500 tokens, providing sufficient headroom above the expected "
    "250-token response (a classification label and 2–4 sentence rationale) while keeping "
    "costs bounded. Because API providers use different parameter names for the same "
    "concepts, a translation mapping is maintained: for example, what OpenAI calls "
    "\u201cmax_tokens\u201d Google calls \u201cmaxOutputTokens.\u201d This mapping ensures "
    "that each API receives its parameters in the format it expects, and that any parameter "
    "unsupported by a given provider is silently excluded from the request."
)

h3("Execution Planning and Session Management")

body(
    "The benchmark produces 5,400 API calls in total: 150 evaluation fragments \u00d7 "
    "6 model families \u00d7 2 prompt conditions \u00d7 3 independent runs per "
    "combination. Because API providers typically impose rate limits, and calling volumes "
    "of this magnitude require careful session management, the batch orchestration system "
    "divides this work into discrete sessions of approximately 1,000 calls each, yielding "
    "roughly six sessions. Within each session, fragments are processed sequentially. For "
    "each fragment, all six models are called under both prompt conditions and all three "
    "runs before the system advances to the next fragment. Between runs within the same "
    "model-prompt combination, a five-second pause is observed to avoid triggering "
    "rate-limit responses from providers. This within-fragment completion-first ordering "
    "ensures that incomplete data are distributed randomly rather than concentrated among "
    "particular fragments if a session is interrupted."
)

h3("Reproducible Execution Ordering")

body(
    "To prevent any systematic relationship between the order in which fragments are "
    "processed and the gold standard classification sequence, the batch orchestrator "
    "applies a reproducible shuffle to the fragment list before processing begins. The "
    "shuffle is seeded with a fixed value (42), meaning that the identical randomized "
    "ordering can be reproduced in any subsequent session or audit. The resulting mapping "
    "of original fragment identifiers to execution positions is logged to a CSV file for "
    "complete audit traceability. A progress-tracking function reads the live response "
    "database and reports the number of successful calls, the number flagged with errors, "
    "and the count of fully complete fragments, enabling the researcher to monitor session "
    "completion without interrupting execution."
)

# ── G.2 ──────────────────────────────────────────────────────────────────────
h2("G.2  Error Handling and Retry Logic")

body_with_link(
    "This module manages the inevitability of transient failures when making thousands of "
    "web-based API requests. A single failed call, if unrecovered, would create a gap in "
    "the dataset that could not easily be repaired. The retry system (",
    "src/api_client/retry.py",
    f"{GITHUB_BASE}/src/api_client/retry.py",
    ") ensures that transient failures are retried automatically while permanent failures "
    "are logged for a separate recovery session."
)

h3("Error Classification")

body(
    "When an API call fails, the system first classifies the type of failure before "
    "deciding how to respond. Six error categories are distinguished. Timeout errors occur "
    "when a provider does not respond within 60 seconds. Rate-limit errors indicate that "
    "the request volume has exceeded the provider\u2019s allowed threshold. "
    "Service-unavailable errors reflect temporary server-side problems at the provider. "
    "Safety-policy triggers occur when a model declines to respond because its content "
    "filters flag the prompt. Invalid-response errors arise when the returned data cannot "
    "be parsed as expected. Generic API errors cover authentication failures, malformed "
    "requests, and other provider-reported problems."
)

h3("Retry Schedule")

body(
    "Transient errors\u2014timeouts, rate limits, and service unavailability\u2014are "
    "treated as retriable, meaning the system will attempt the call again after a waiting "
    "period. The waiting schedule follows an exponential pattern: the first retry waits "
    "10 seconds, the second waits 30 seconds, and the third waits 90 seconds. This "
    "graduated approach gives providers time to recover without flooding them with rapid "
    "repeated requests. Each call is allowed a maximum of three total attempts (one initial "
    "call plus two retries). Permanent errors, such as safety triggers and authentication "
    "failures, are not retried, because additional attempts would not resolve the underlying "
    "problem. For unrecognized error types, a single retry is permitted. After all retry "
    "opportunities are exhausted, the system records the failure and moves to the next call."
)

h3("Failed Call Logging and Recovery")

body(
    "Failed calls are persisted to a JSON-Lines log file\u2014one JSON record per "
    "line\u2014so that they can be re-attempted in a dedicated recovery session rather "
    "than leaving the primary data collection workflow in an indefinite loop. The log "
    "captures the fragment identifier, model, prompt condition, run number, error category, "
    "and error message. A corresponding loading function reads this log at the start of a "
    "recovery session, and a rerun function re-submits each failed call through the same "
    "retry pipeline with a five-second pause between calls. After a successful recovery "
    "session, the log is cleared to prevent accidental re-processing in future sessions. "
    "This design ensures that no call is permanently lost due to transient infrastructure "
    "failures, and that the final response database is as complete as possible."
)

# ── G.3 ──────────────────────────────────────────────────────────────────────
h2("G.3  Request Construction and Response Capture")

p = doc.add_paragraph(
    "This module handles the precise mechanics of building an API request, sending it, "
    "receiving the response, extracting the relevant content, and immediately saving a "
    "structured record to the response database. It also normalizes and validates model "
    "outputs before they are stored. The relevant source files are "
)
p.style = doc.styles["Normal"]
p.paragraph_format.space_after = Pt(6)
add_hyperlink(p, "src/api_client/executor.py", f"{GITHUB_BASE}/src/api_client/executor.py")
p.add_run(" and ")
add_hyperlink(p, "src/api_client/parser.py", f"{GITHUB_BASE}/src/api_client/parser.py")
p.add_run(".")

h3("Request Construction")

body(
    "Before any call is made, the executor loads the appropriate prompt template from the "
    "configuration directory, inserting the evaluation fragment text into the designated "
    "placeholder position. The zero-shot template contains the task description and "
    "instructions; the few-shot template additionally includes four annotated calibration "
    "examples (two classified as sound and two as not sound) before the target fragment. "
    "Once the prompt is rendered, HTTP authentication headers are constructed for the "
    "target provider. The request body is assembled in the format each provider expects: "
    "OpenAI-compatible providers receive a \u201cmessages\u201d array with a single user "
    "turn; Anthropic\u2019s API requires that the maximum token count appear at the top "
    "level of the request rather than nested within parameters; Google\u2019s Gemini API "
    "expects a different content structure with nested sub-fields. Parameter filtering "
    "ensures that unsupported options are excluded for each provider. API latency is "
    "measured precisely for each call as the elapsed time between sending the request and "
    "receiving the response, and this value is stored with the response record."
)

h3("Response Capture and Parsing")

body(
    "After a successful API call, the returned response is parsed in two stages. First, "
    "the raw text content is extracted from the provider-specific response envelope. Three "
    "response formats are handled: OpenAI-compatible responses nest the text in a "
    "\u201cchoices\u201d array; Anthropic responses nest it in a \u201ccontent\u201d "
    "array; Gemini responses nest it in a \u201ccandidates\u201d array. Second, the "
    "extracted text is parsed to recover the model\u2019s classification (\u201csound\u201d "
    "or \u201cnot sound\u201d) and its written rationale. The parser attempts structured "
    "JSON extraction first, stripping any markdown code fences that models sometimes wrap "
    "around their output. If JSON parsing fails\u2014because a model returned free text "
    "instead\u2014the parser falls back to pattern matching that searches for "
    "classification keywords and labeled \u201crationale:\u201d fields. The phrase "
    "\u201cnot sound\u201d is always searched before the bare word \u201csound\u201d "
    "to prevent false matches where \u201csound\u201d appears as a substring of "
    "\u201cnot sound.\u201d"
)

h3("Classification Normalization and Rationale Validation")

body(
    "Raw classification strings are normalized to one of two canonical values: "
    "\u201csound\u201d or \u201cnot_sound.\u201d The normalization routine handles "
    "common variant phrasings\u2014for example, \u201cvalid,\u201d \u201cacceptable,\u201d "
    "and \u201cdefensible\u201d are recognized as sound; \u201cunsound,\u201d "
    "\u201cinvalid,\u201d and \u201cweak\u201d are recognized as not sound\u2014again "
    "prioritizing not-sound variants in the matching sequence to prevent erroneous sound "
    "matches. A rationale is considered valid if it contains at least ten words and "
    "includes at least two keywords from a controlled list of domain-reasoning terms such "
    "as \u201cevidence,\u201d \u201csynthesis,\u201d \u201cjudgment,\u201d or "
    "\u201canalysis.\u201d This threshold ensures that a rationale contains substantive "
    "reasoning rather than merely restating the classification label. Responses that fail "
    "either test receive an error flag in the response database, and the specific problem "
    "is recorded in an error-details field."
)

h3("Immediate Persistence")

body(
    "Each response record is appended to the response database CSV file immediately after "
    "processing, rather than accumulated in memory and saved at session end. This design "
    "ensures that a session interruption or system failure can lose at most a single "
    "in-flight call rather than an entire session\u2019s worth of data. The database "
    "schema includes a unique response identifier, fragment and model identifiers, the "
    "prompt condition and run number, the normalized classification and rationale text, "
    "timestamps, API latency in seconds, input and output token counts, the API version "
    "string, the parsing method used, and the error flag with details."
)

# ── G.4 ──────────────────────────────────────────────────────────────────────
h2("G.4  Internal Coherence Validation")

p = doc.add_paragraph(
    "This module implements the first substantive step of the automated scoring pipeline: "
    "assessing whether each model response demonstrates internal logical consistency. A "
    "response is considered internally coherent if the written rationale logically supports "
    "the stated classification. A model that classifies a fragment as \u201csound\u201d "
    "while writing a rationale that describes weaknesses has produced an internally "
    "inconsistent response, and no amount of accidental correctness against the gold "
    "standard can make such a response analytically meaningful. The relevant source files "
    "are "
)
p.style = doc.styles["Normal"]
p.paragraph_format.space_after = Pt(6)
add_hyperlink(p, "src/scoring/coherence.py", f"{GITHUB_BASE}/src/scoring/coherence.py")
p.add_run(" and ")
add_hyperlink(p, "src/scoring/config.py", f"{GITHUB_BASE}/src/scoring/config.py")
p.add_run(".")

h3("Tiered Validation Design")

body(
    "Coherence validation employs a two-tier architecture that balances speed, cost, and "
    "accuracy. The majority of responses can be classified confidently using the faster, "
    "cost-free Tier 1 method; only genuinely ambiguous cases require the more expensive "
    "Tier 2 LLM screen."
)

h3("Tier 1: Rule-Based Keyword Matching")

body(
    "The first tier counts how many \u201cstrength indicator\u201d words and how many "
    "\u201cweakness indicator\u201d words appear in the rationale text. The strength "
    "indicator list contains adjectives and verbs associated with positive evaluation "
    "quality\u2014words such as \u201cclear,\u201d \u201crobust,\u201d "
    "\u201cdemonstrates,\u201d \u201cintegrates,\u201d and \u201cadequately.\u201d "
    "The weakness indicator list contains terms associated with identified "
    "problems\u2014words such as \u201cmissing,\u201d \u201cinsufficient,\u201d "
    "\u201cfails to,\u201d \u201cflawed,\u201d and \u201ccontradictory.\u201d A complete "
    "list of both indicator sets is maintained in the scoring configuration file."
)

body(
    "The coherence signal from Tier 1 is determined by the relationship between the "
    "classification and the dominant indicator type. If the model classified the fragment "
    "as sound and strength indicators outnumber weakness indicators, the rationale is "
    "flagged as coherent. If the model classified the fragment as not sound and weakness "
    "indicators outnumber strength indicators, the rationale is also flagged as coherent. "
    "Any mismatch\u2014a \u201csound\u201d classification with weakness-dominated "
    "language, or a \u201cnot sound\u201d classification with strength-dominated "
    "language\u2014produces an incoherent signal."
)

body(
    "The system is considered confident in its Tier 1 judgment when the absolute "
    "difference between the strength and weakness indicator counts is at least three "
    "(the confidence threshold). When this threshold is met, the Tier 1 judgment is "
    "accepted as the final coherence determination and no further processing is performed. "
    "When the difference falls below this threshold\u2014indicating the rationale contains "
    "a mixture of evaluative language that does not clearly favor one pole\u2014the "
    "response is routed to Tier 2."
)

h3("Tier 2: LLM Semantic Screening")

body(
    "The second tier submits the response to Claude Haiku (a lightweight Anthropic model) "
    "at temperature zero, asking it to judge whether the rationale logically supports the "
    "classification. Critically, the Tier 2 prompt provides only the classification and "
    "the rationale text\u2014it does not include the evaluation fragment being assessed, "
    "the gold standard classification, or any information about the study. This design "
    "preserves the independence of the coherence check from the benchmarked task, "
    "preventing circular reasoning in which the screening model\u2019s knowledge of the "
    "\u201ccorrect\u201d answer could influence its coherence judgment. The prompt "
    "instructs the screening model to respond with a single word: \u201ccoherent\u201d "
    "or \u201cincoherent.\u201d Because the word \u201cincoherent\u201d contains "
    "\u201ccoherent\u201d as a substring, the parsing logic checks for "
    "\u201cincoherent\u201d first before checking for the shorter string."
)

h3("Tier Agreement and Ambiguous Cases")

body(
    "If Tier 1 and Tier 2 agree on a coherence judgment, that judgment is recorded as "
    "final regardless of direction. If they disagree\u2014a situation expected to be "
    "rare\u2014the response is flagged as \u201cambiguous\u201d and marked for manual "
    "researcher review. Responses that failed the earlier parsing or rationale-validation "
    "steps are assigned \u201cincoherent\u201d directly, without invoking either tier, "
    "since they already represent incomplete or unacceptable outputs."
)

h3("RC2 Reliability Check")

body(
    "To validate that the automated coherence pipeline agrees acceptably with human "
    "judgment, the module supports an RC2 reliability check. A stratified sample of "
    "responses is selected for manual review: 20 classified as coherent by the pipeline, "
    "20 classified as incoherent, and up to 10 flagged as ambiguous. The researcher "
    "independently judges each sampled response as coherent or incoherent. Cohen\u2019s "
    "kappa is then calculated between the pipeline\u2019s final determinations and the "
    "researcher\u2019s manual judgments. The study targets a kappa of at least 0.80; "
    "if this threshold is not met, the coherence pipeline must be recalibrated before "
    "analysis proceeds."
)

# ── G.5 ──────────────────────────────────────────────────────────────────────
h2("G.5  Scoring Pipeline and Accuracy Algorithms")

p = doc.add_paragraph(
    "This module implements the remaining three steps of the four-step automated scoring "
    "pipeline and provides the functions that make the full pipeline runnable as a single "
    "coordinated operation. The relevant source files are "
)
p.style = doc.styles["Normal"]
p.paragraph_format.space_after = Pt(6)
add_hyperlink(p, "src/scoring/accuracy.py", f"{GITHUB_BASE}/src/scoring/accuracy.py")
p.add_run(", ")
add_hyperlink(p, "src/scoring/edge_cases.py", f"{GITHUB_BASE}/src/scoring/edge_cases.py")
p.add_run(", and ")
add_hyperlink(p, "src/scoring/pipeline.py", f"{GITHUB_BASE}/src/scoring/pipeline.py")
p.add_run(".")

h3("Pre-Step: Edge Case Detection")

body(
    "Before coherence validation begins, each response passes through an edge case "
    "detection screen. This pre-step categorizes responses that fall outside the normal "
    "processing path and determines how they should be handled. The detection logic "
    "proceeds through an ordered priority sequence. Technical timeouts are identified "
    "first and are excluded from the analysis entirely, because a timeout represents an "
    "infrastructure failure rather than a model reasoning failure. Safety-policy triggers "
    "are flagged separately. Substantive refusals\u2014responses where the model explicitly "
    "declines to evaluate the fragment using phrases such as \u201ccannot evaluate\u201d "
    "or \u201cinsufficient information\u201d\u2014are scored as failures. Responses with "
    "missing classifications or rationales shorter than ten words are scored as failures. "
    "Rationales containing an excessively high proportion of repeated words (below thirty "
    "percent unique words) are flagged as nonsensical. Responses that hedge without "
    "reaching a conclusion are checked for a clear final classification using conclusory "
    "markers such as \u201ctherefore\u201d or \u201cin conclusion.\u201d If no such marker "
    "is present, the response is scored as a failure due to unresolved ambivalence. "
    "Responses that assert both \u201csound\u201d and \u201cnot sound\u201d are detected "
    "using word-boundary regular expressions that correctly distinguish the phrase "
    "\u201cnot sound\u201d from standalone occurrences of \u201csound.\u201d"
)

h3("Step 2: Classification Accuracy")

body(
    "Once coherence is established for each response, coherent responses are compared to "
    "the locked gold standard expert classifications. The gold standard file is identified "
    "by a date-stamped naming convention that prevents ambiguity about which file is "
    "authoritative. Only responses where the automated pipeline determined the rationale "
    "to be coherent are evaluated for accuracy. Incoherent responses automatically receive "
    "a failing outcome regardless of whether their classification happens to match the "
    "gold standard, because an incoherent response cannot be treated as a reliable "
    "judgment. The accuracy field is set to \u201ccorrect\u201d when the normalized model "
    "classification matches the expert classification and to \u201cincorrect\u201d "
    "otherwise."
)

h3("Step 3: Run-Level Outcome Assignment")

body(
    "Each of the 5,400 individual runs receives a binary Pass or Fail outcome based on "
    "the combination of its coherence status and its classification accuracy. A run passes "
    "only if it is both internally coherent and correct against the gold standard. Any "
    "other combination constitutes a failure. Failed runs are additionally classified into "
    "one of three failure types. Type 1 failures\u2014coherent but incorrect\u2014represent "
    "the analytically most significant category: the model maintained logical consistency "
    "in its reasoning but arrived at the wrong conclusion. Type 1 failures are the only "
    "category eligible for the detailed failure mode coding analysis described in Appendix "
    "H.3. Type 2 failures (incoherent but accidentally matching the gold standard) and "
    "Type 3 failures (incoherent and incorrect) are documented quantitatively but receive "
    "no further qualitative analysis."
)

h3("Step 4: Fragment-Level Adjudication")

body(
    "The three independent runs for each fragment-model-prompt combination are adjudicated "
    "using a majority rule: a combination passes if at least two of its three runs pass. "
    "A combination fails if at least two of its three runs fail. The system additionally "
    "tracks whether all three runs agree\u2014unanimous passes and unanimous failures are "
    "noted separately as indicators of response consistency. The result is a dataset of "
    "1,800 adjudicated outcomes (150 fragments \u00d7 6 models \u00d7 2 prompt "
    "conditions), each representing the definitive judgment for that combination. These "
    "adjudicated outcomes serve as the primary unit of analysis for the statistical tests "
    "in Appendix H.2. Fragment-level outcomes are saved to a separate CSV file, and "
    "performance metrics\u2014overall pass rates, unanimous agreement rates, and "
    "stratifications by model and prompt condition\u2014are computed and exported at "
    "this stage. For the primary diagnostic model (GPT 5.2), fragment-level Type 1 "
    "failures are identified and exported to a separate file that serves as the input "
    "for the manual failure mode coding described in Appendix H.3."
)

# ─────────────────────────────────────────────────────────────────────────────
# APPENDIX H
# ─────────────────────────────────────────────────────────────────────────────

doc.add_page_break()

h1("Appendix H: Analysis Infrastructure")

body(
    "This appendix documents the software modules used to analyze the scored dataset "
    "produced by the pipeline described in Appendix G. Five subsections correspond to "
    "the five primary analysis activities: generating the performance matrix (H.1), "
    "running the pre-specified statistical tests (H.2), synthesizing GPT 5.2 failure "
    "patterns from the manually coded failure dataset (H.3), identifying challenge cases "
    "in the fragment corpus (H.4), and conducting the meta-evaluation and validity "
    "assessment (H.5)."
)

# ── H.1 ──────────────────────────────────────────────────────────────────────
h2("H.1  Performance Matrix Generation")

body_with_link(
    "This module assembles and exports the primary performance tables that form the "
    "quantitative foundation of the study\u2019s findings. The relevant source file is ",
    "src/analysis/performance.py",
    f"{GITHUB_BASE}/src/analysis/performance.py",
    "."
)

h3("The 12-Cell Performance Matrix")

body(
    "The central output of this module is a performance matrix presenting fragment-level "
    "pass rates for each of the twelve model-prompt combinations (six model families "
    "\u00d7 two prompt conditions). Each cell in the matrix reports the observed pass "
    "rate along with a 95% confidence interval calculated using the Wilson score method. "
    "The Wilson score interval is preferred over the more familiar plus-or-minus "
    "interval because it produces accurate coverage even for small samples and extreme "
    "proportions\u2014situations where a model achieves very high or very low pass rates. "
    "The interval bounds are clamped to the permissible range of zero to one to prevent "
    "mathematically possible but logically impossible values such as a negative lower "
    "bound. The matrix is generated in a fixed canonical order (closed models first, then "
    "open models; zero-shot before few-shot) regardless of the order in which data happen "
    "to appear in the underlying dataset, ensuring that output tables are consistent "
    "across runs."
)

h3("Collapsed Summaries")

body(
    "Alongside the full matrix, three collapsed summary tables are generated. The first "
    "collapses across prompt conditions, reporting each model\u2019s average pass rate "
    "across both zero-shot and few-shot conditions. The second collapses across models, "
    "reporting the average pass rate for zero-shot and few-shot conditions independently. "
    "The third collapses by architectural category\u2014open-source models versus closed "
    "proprietary models\u2014reporting aggregate pass rates for each architecture type "
    "with confidence intervals. The architecture classification (which models are "
    "\u201copen\u201d and which are \u201cclosed\u201d) is maintained centrally in the "
    "configuration file rather than defined locally within each analysis function, "
    "ensuring consistency across all modules that reference this distinction. All four "
    "tables are exported as CSV files to the results directory. Confidence interval widths "
    "are included in the full matrix table to provide a direct indicator of estimation "
    "precision for each cell."
)

# ── H.2 ──────────────────────────────────────────────────────────────────────
h2("H.2  Statistical Tests")

body_with_link(
    "This module implements the three inferential statistical tests specified in the "
    "study protocol. The relevant source file is ",
    "src/analysis/statistics.py",
    f"{GITHUB_BASE}/src/analysis/statistics.py",
    "."
)

h3("Test 1: McNemar\u2019s Test for Prompt Condition Effects")

body(
    "The first test evaluates whether the prompt condition\u2014zero-shot versus "
    "few-shot\u2014produces a statistically significant difference in pass rates. Because "
    "each evaluation fragment is assessed under both conditions, the data have a natural "
    "paired structure: for every fragment-model combination, there exists a zero-shot "
    "outcome and a few-shot outcome. McNemar\u2019s test is designed precisely for paired "
    "binary data of this kind. It focuses on \u201cdiscordant\u201d pairs\u2014those "
    "where the two conditions produced different outcomes\u2014and asks whether the "
    "discordance is systematically tilted in one direction (for example, fragments that "
    "pass under few-shot but fail under zero-shot appear significantly more often than "
    "the reverse)."
)

body(
    "When the number of discordant pairs is small (fewer than 25), the test uses an "
    "exact binomial probability calculation, which is more accurate than the approximation "
    "used for larger samples. When 25 or more discordant pairs exist, a chi-squared "
    "approximation with continuity correction is applied. Both variants produce a p-value "
    "that is compared against the 0.05 significance threshold. The test also calculates "
    "an odds ratio\u2014the ratio of discordant pairs in each direction\u2014which "
    "provides a readily interpretable effect size. An odds ratio greater than one indicates "
    "an advantage for few-shot over zero-shot; an odds ratio less than one indicates an "
    "advantage for zero-shot. Confidence intervals for the odds ratio are computed using "
    "the log method when both discordant cell counts are nonzero. The overall test pools "
    "all 900 fragment-model pairs (150 fragments \u00d7 6 models), providing the primary "
    "inferential result. Per-model breakdowns are additionally computed to provide "
    "diagnostic insight into whether the prompt-condition effect is uniform across all "
    "models or concentrated in particular architectures."
)

h3("Test 2: Chi-Squared Test for Model Family Differences")

body(
    "The second test evaluates whether the six model families differ significantly in "
    "their overall pass rates. A contingency table crossing model family against pass/fail "
    "outcome is constructed from the fragment-level outcomes, and a chi-squared test of "
    "independence is applied to this table. The omnibus test assesses overall significance. "
    "If the omnibus test reaches significance, pairwise comparisons are conducted for all "
    "fifteen model-pair combinations, with a Bonferroni correction applied to control the "
    "family-wise error rate across multiple comparisons. Effect size is quantified using "
    "Cram\u00e9r\u2019s V, a measure that ranges from zero (no association) to one "
    "(perfect association) and which accounts for the size of the contingency table."
)

h3("Test 3: Z-Test for Architecture Comparison")

body(
    "The third test compares the aggregate pass rates of closed proprietary models (GPT "
    "5.2, Claude Opus 4.6, and Gemini 3 Pro) against open-source models (DeepSeek V3, "
    "Kimi K2, and GLM 4.7). This is a two-group comparison of independent proportions, "
    "tested using a z-statistic based on the pooled proportion under the null hypothesis "
    "of no difference. Effect size is quantified using Cohen\u2019s h, a measure "
    "specifically designed for proportions that is analogous to Cohen\u2019s d for means. "
    "Values below 0.2 indicate small effects, values between 0.2 and 0.5 indicate medium "
    "effects, and values above 0.5 indicate large effects. All three test results are "
    "exported as a structured JSON file to the results directory."
)

# ── H.3 ──────────────────────────────────────────────────────────────────────
h2("H.3  Failure Pattern Synthesis for GPT 5.2")

body_with_link(
    "This module synthesizes the manually coded failure modes for the primary diagnostic "
    "model (GPT 5.2) into quantitative patterns. It operates on a dataset of failure "
    "codes assigned by the researcher during the Phase 3 qualitative analysis, using the "
    "six-domain Integrated Evaluative Reasoning Checklist framework. The relevant source "
    "file is ",
    "src/analysis/failure_patterns.py",
    f"{GITHUB_BASE}/src/analysis/failure_patterns.py",
    "."
)

h3("Within-Domain Analysis")

body(
    "For each of the six evaluative reasoning domains\u2014Evaluative Framing, Evidence "
    "Base, Argument Structure, Synthesis and Integration, Evaluative Conclusion, and "
    "Qualifications and Transparency\u2014the module calculates the count and percentage "
    "of GPT 5.2 Type 1 failures whose primary code belongs to that domain. It also "
    "identifies the most frequently violated specific checkpoints within each domain and "
    "computes the number of failures in each domain that occurred under zero-shot versus "
    "few-shot conditions, yielding a \u201cprompt differential\u201d that indicates "
    "whether the domain\u2019s failures are sensitive to the calibration examples or "
    "appear to be prompt-resistant. Domain codes are identified using word-boundary "
    "anchored pattern matching to prevent false matches against non-standard or emergent "
    "code strings."
)

h3("Cross-Domain Co-occurrence")

body(
    "Because each coded failure case may receive up to three codes (a primary code and "
    "optionally a secondary and tertiary code reflecting compounding problems), the module "
    "builds a co-occurrence matrix counting how often failures from each pair of domains "
    "appear together in the same case. This analysis extends beyond the primary code "
    "alone\u2014when two domains co-occur frequently, it suggests that failures in those "
    "areas compound each other and may share an underlying cause. All pairwise "
    "combinations of up to three codes per case are considered, ensuring that compounding "
    "failures captured in secondary and tertiary codes are not overlooked."
)

h3("Emergent Code Analysis")

body(
    "The failure coding protocol permits the researcher to generate new \u201cinductive\u201d "
    "codes for failure patterns not anticipated by the pre-specified framework, alongside "
    "the expected \u201cdeductive\u201d codes derived from the checklist. The module "
    "quantifies the proportion of failures receiving inductive codes. If fewer than 20% "
    "of failures require inductive codes, the framework is assessed as adequately covering "
    "the observed failure space. Higher proportions indicate that the framework required "
    "meaningful supplementation beyond its original design."
)

h3("Diagnostic Profile")

body(
    "The module assembles a structured diagnostic summary of GPT 5.2\u2019s failure "
    "profile: the two domains in which it fails most frequently, the five most common "
    "specific checkpoint violations, the domains that appear particularly sensitive to "
    "the prompt condition (those where the zero-shot and few-shot failure counts differ "
    "by more than three), and the domains that appear resistant to prompt condition (those "
    "where the counts differ by at most one despite a meaningful number of failures, "
    "suggesting the failures persist regardless of calibration). These elements are "
    "exported as CSV and JSON files to the results directory."
)

# ── H.4 ──────────────────────────────────────────────────────────────────────
h2("H.4  Challenge Case Identification")

body_with_link(
    "This module identifies fragments that warrant special analytical attention as "
    "\u201cchallenge cases\u201d\u2014those that prove systematically difficult across "
    "multiple dimensions of the study. The relevant source file is ",
    "src/analysis/challenge_cases.py",
    f"{GITHUB_BASE}/src/analysis/challenge_cases.py",
    "."
)

h3("Criterion 1: Systematic Disagreement")

body(
    "The first challenge case criterion identifies fragment-prompt combinations where GPT "
    "5.2 failed and at least two of the five comparison models also failed on the same "
    "combination. Fragments meeting this criterion represent cases where multiple models "
    "across different architectures and training regimes struggled, suggesting that the "
    "difficulty originates in the evaluation fragment itself rather than in idiosyncratic "
    "weaknesses of a single model. A precomputed lookup table is used to match GPT 5.2 "
    "failures against comparison model failures efficiently, avoiding the computational "
    "cost of repeated full-dataset searches that would otherwise make this analysis slow "
    "on large datasets. Fragment-prompt identifiers are represented as structured tuples "
    "rather than concatenated text strings to prevent accidental false matches that could "
    "arise if underscores appear in fragment identifiers."
)

h3("Criterion 2: Recurring Failure Mode")

body(
    "The second challenge case criterion identifies specific failure codes that account "
    "for at least ten percent of all GPT 5.2 Type 1 failures. A failure code meeting this "
    "threshold represents a systematic reasoning weakness that recurs across multiple "
    "fragments rather than being unique to individual cases. When a code accounts for "
    "this proportion of total failures, the fragments it affects collectively constitute "
    "a distinct challenge zone in the fragment corpus\u2014a category of evaluation task "
    "that GPT 5.2 consistently struggles to perform correctly."
)

h3("Overlap Analysis")

body(
    "After applying both criteria independently, the module examines the overlap between "
    "the sets of fragment-prompt combinations they identify. High overlap\u2014where most "
    "combinations appear under both criteria\u2014suggests that challenge cases reflect "
    "genuinely difficult fragments that are hard for models to reason about correctly "
    "across multiple dimensions. Low overlap suggests that the two criteria capture "
    "different aspects of difficulty: Criterion 1 may identify fragments with inherently "
    "ambiguous evaluative content that challenges all models broadly, while Criterion 2 "
    "may identify fragments that specifically expose a systematic gap in GPT 5.2\u2019s "
    "reasoning repertoire. The overlap analysis and summary tables are exported to the "
    "results directory."
)

# ── H.5 ──────────────────────────────────────────────────────────────────────
h2("H.5  Meta-Evaluation and Validity Assessment")

body_with_link(
    "This module provides the structured framework for assessing the validity of the "
    "study\u2019s conclusions. It synthesizes reliability check results, evaluates the "
    "defensibility of the gold standard, organizes the construct validity evidence, and "
    "catalogs pre-specified methodological limitations. The relevant source file is ",
    "src/analysis/meta_evaluation.py",
    f"{GITHUB_BASE}/src/analysis/meta_evaluation.py",
    "."
)

h3("Construct Validity Triangulation")

body(
    "Three structured questions are posed to triangulate the construct validity of the "
    "evaluative reasoning measure. The first asks whether GPT 5.2 struggles more with "
    "Synthesis and Integration (Domain 4) than with Evidence Identification (Domain 2), "
    "which would be consistent with theoretical expectations that higher-order integration "
    "tasks are harder for current language models than recognizing the presence of "
    "evidence. The second asks whether few-shot calibration provides proportionally greater "
    "benefit for domains that rely on tacit evaluative judgment (such as Domain 6, "
    "Qualifications and Transparency), where calibration examples might communicate "
    "implicit standards that cannot be articulated as explicit rules. The third asks "
    "whether challenge cases cluster around evaluation criteria requiring complex "
    "counterfactual reasoning, as would be expected if fragments testing sustainability "
    "or long-term adequacy judgments are genuinely harder to evaluate correctly. These "
    "questions serve as structured prompts for post-data-collection interpretation; the "
    "code scaffolds the collection of findings but does not compute them automatically."
)

h3("Gold Standard Defensibility")

body(
    "The module identifies which fragments were flagged as challenge cases under Criterion "
    "1 of the challenge case analysis and notes these as candidates for analytical "
    "reflection against the gold standard. The gold standard itself remains locked and "
    "cannot be revised; this section documents the researcher\u2019s reflective engagement "
    "with any systematic model disagreements, providing a template for recording whether "
    "the original classification was clearly defensible, whether the models revealed "
    "genuine ambiguity in the fragment, or whether the disagreement represents an "
    "acknowledged limitation. The count of boundary cases originally flagged in the gold "
    "standard during the expert classification phase is also surfaced here."
)

h3("Reliability Synthesis")

body(
    "Three reliability checks are synthesized in this module. RC1 assesses the expert\u2019s "
    "temporal consistency: the same fragments are re-classified a minimum of two weeks "
    "after the original classification, and Cohen\u2019s kappa is calculated between the "
    "two sets of classifications. RC2 assesses the automated coherence pipeline\u2019s "
    "agreement with manual human judgment, as described in Appendix G.4. RC3 assesses the "
    "researcher\u2019s consistency in applying the failure mode coding protocol by "
    "re-coding a sample of previously coded cases after an interval. All three checks "
    "target a kappa of at least 0.80. The module handles reliability result dictionaries "
    "in two formats\u2014a simple flat structure and a nested structure\u2014to "
    "accommodate flexibility in how results are passed from the reliability-checking "
    "functions. The overall validity verdict is derived only when all three checks have "
    "been completed; if any remain pending at the time the module is run, the overall "
    "assessment is withheld pending completion."
)

h3("Methodological Limitations")

body(
    "Five pre-specified methodological limitations are cataloged with their associated "
    "threats, mitigations, and residual risks. The limitations concern (1) the "
    "representativeness of conclusion sections as samples of evaluative reasoning, "
    "(2) the corpus scope\u2019s restriction to UN evaluation reports, (3) the reliance "
    "on a single expert for gold standard classifications rather than a consensus panel, "
    "(4) the use of a fixed prompt formulation per condition rather than optimized prompt "
    "engineering, and (5) the concentration of detailed failure mode coding on the single "
    "primary diagnostic model. The module exports both a structured JSON version and a "
    "readable Markdown narrative version of the full meta-evaluation report to the results "
    "directory."
)

# ── Save ─────────────────────────────────────────────────────────────────────
OUTPUT = "/Users/zachtilton/dissertation/eval-reasoning-benchmark/appendices/Appendices_G_and_H.docx"
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
